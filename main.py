import os
import csv
from time import sleep
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

INPUT_FILE = "200. How People Kept Warm in Castles Without Fireplaces ENG.docx"
OUTPUT_FILE = "video_prompts.csv"
MAX_PARAGRAPHS_TO_PROCESS = 5

SYSTEM_PROMPT = (
    "You are a film director, anthropologist and visual historian. "
    "Create one concise cinematic video prompt in English "
    "for Google Veo 3 (fast mode) based on the given paragraph."
)


def extract_numbered_paragraphs(doc_path):
    """
    Извлекает все абзацы с текстом из DOCX-файла.
    """
    paragraphs = []

    try:
        document = Document(doc_path)
        print(f"Загружен, всего параграфов: {len(document.paragraphs)}")
    except Exception as e:
        print(f"Не получается открыть файл: {e}")
        return paragraphs
    para_counter = 1
    for para in document.paragraphs:
        text = para.text.strip()
        # Пропускаем пустые строки и очень короткий текст
        if not text or len(text) < 10:
            continue
        # Пропускаем заголовки глав
        if text.startswith("Chapter ") or text.startswith("**Chapter"):
            continue
        # Пропускаем заголовок "Introduction" и "Conclusion"
        if text in ["Introduction", "Conclusion"]:
            continue
        # Добавляем абзац с искусственной нумерацией
        paragraphs.append({
            "number": str(para_counter),
            "text": text
        })
        para_counter += 1
    print(f"Извлечено абзацев: {len(paragraphs)}")
    # Показываем первые 3 для проверки
    for i, p in enumerate(paragraphs[:3]):
        print(f"  Абзац {p['number']}: {p['text'][:60]}...")
    return paragraphs


def save_results(results, path):
    """
    Сохраняет сгенерированные видео-промпты в CSV-файл.
    """
    if not results:
        print("Нечего сохранять, результаты пустые")
        return False
    try:
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(
                f,
                fieldnames=["Paragraph #", "Script paragraph", "Video prompt"]
            )
            writer.writeheader()
            writer.writerows(results)
        return True
    except Exception as e:
        print(f"Ошибка при сохранении файла {path}: {e}")
        return False


def generate_video_prompt(client, paragraph_text, retries=3):
    """
    Отправляет абзац сценария в ChatGPT и получает видео-промпт.
    """
    if not paragraph_text or len(paragraph_text.strip()) < 10:
        return None

    for attempt in range(retries):
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {
                        "role": "user",
                        "content": f"Create a video prompt:\n{paragraph_text}"
                    }
                ],
                temperature=0.8,
                max_tokens=120,
            )
            prompt = response.choices[0].message.content.strip()
            return prompt if prompt else None

        except Exception as e:
            if attempt < retries - 1:
                wait = 2 ** attempt
                print(f"API ошибка, повтор через {wait} секунд ({e})")
                sleep(wait)
            else:
                print(f"Не удалось получить ответ {retries} попыток: {e}")
                return None


def main():
    """
    Основная точка входа в скрипт.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("Ошибка: переменная OPENAI_API_KEY не установлена.")
        print("Windows: set OPENAI_API_KEY=your_key_here")
        print("Linux/Mac: export OPENAI_API_KEY=your_key_here")
        return

    if not os.path.exists(INPUT_FILE):
        print(f"Файл не найден: {INPUT_FILE}")
        return

    print("=" * 50)
    print("Запуск обработки сценария")
    print("=" * 50)

    client = OpenAI(api_key=api_key)

    paragraphs = extract_numbered_paragraphs(INPUT_FILE)
    if not paragraphs:
        print("Не найдено нумерованных абзацев в документе")
        return

    print(f"Найдено абзацев: {len(paragraphs)}")

    paragraphs_to_process = (
        paragraphs[:MAX_PARAGRAPHS_TO_PROCESS]
        if MAX_PARAGRAPHS_TO_PROCESS
        else paragraphs
    )

    results = []

    for idx, para in enumerate(paragraphs_to_process, 1):
        print(f"\n[{idx}/{len(paragraphs_to_process)}] Абзац {para['number']}")

        prompt = generate_video_prompt(client, para["text"])
        if prompt:
            results.append({
                "Paragraph #": para["number"],
                "Script paragraph": para["text"],
                "Video prompt": prompt
            })
            save_results(results, OUTPUT_FILE)
            print("Промпт успешно получен")
        else:
            print("Не удалось сгенерировать промпт")

        sleep(0.5)

    print("\n" + "=" * 50)
    print(f"Сгенерировано промптов: {len(results)}")
    print(f"Файл результата: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
